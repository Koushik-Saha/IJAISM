import docx
import os
import json
import re

def extract_metadata(file_path):
    doc = docx.Document(file_path)
    metadata = {
        "file_path": file_path,
        "filename": os.path.basename(file_path),
        "title": "",
        "authors": [],
        "emails": [],
        "affiliations": [],
        "abstract": "",
        "keywords": [],
        "article_type": "Research Article",
        "doi": ""
    }

    # Extract all text from paragraphs and tables for better searching
    all_paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_texts = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text: table_texts.append(text)
    
    all_text = "\n".join(all_paragraphs + table_texts)

    # Simple heuristic-based extraction
    ptr = 0
    if all_paragraphs and "Article" in all_paragraphs[ptr]:
        metadata["article_type"] = all_paragraphs[ptr]
        ptr += 1
    
    if ptr < len(all_paragraphs):
        metadata["title"] = all_paragraphs[ptr]
        ptr += 1
        
    if ptr < len(all_paragraphs):
        authors_text = all_paragraphs[ptr]
        clean_authors = re.sub(r'[\d\*]+', '', authors_text)
        # De-duplicate authors and filter empty
        metadata["authors"] = list(dict.fromkeys([a.strip() for a in clean_authors.split(',') if a.strip()]))
        ptr += 1

    # Extract affiliations and emails
    for i in range(ptr, min(ptr + 10, len(all_paragraphs))):
        p = all_paragraphs[i]
        if "Corresponding Author" in p or "@" in p:
            emails = re.findall(r'[\w\.-]+@[\w\.-]+', p)
            metadata["emails"].extend(list(set(emails)))
        elif re.match(r'^\d', p) or "Department" in p or "University" in p:
            metadata["affiliations"].append(p)
            
    # Extract DOI
    doi_match = re.search(r'DOI[:\s]*(https?://doi\.org/\S+|10\.\d{4,9}/\S+)', all_text, re.IGNORECASE)
    if doi_match:
        metadata["doi"] = doi_match.group(1).strip()
    
    # Extract Abstract
    abstract_match = re.search(r'Abstract[:\s]*(.*?)(?:\n\d+\.|\nKeywords|$|DOI:)', all_text, re.DOTALL | re.IGNORECASE)
    if abstract_match:
        metadata["abstract"] = abstract_match.group(1).strip()
    
    # Extract Keywords
    keywords_match = re.search(r'Keywords[:\s]*(.*?)(?:\n|$)', all_text, re.IGNORECASE)
    if keywords_match:
        kw_text = keywords_match.group(1).strip()
        metadata["keywords"] = [k.strip() for k in re.split(r'[,;]', kw_text)]

    return metadata

def main():
    base_dir = "/Users/koushiksaha/Downloads/C5K_V1_I1"
    all_metadata = []
    
    for root, dirs, files in os.walk(base_dir):
        # Skip top level folder as it's not a journal
        if root == base_dir:
            continue
            
        for f in files:
            if f.endswith(".docx") and not f.startswith("~$"):
                path = os.path.join(root, f)
                try:
                    meta = extract_metadata(path)
                    if meta:
                        meta["journal_folder"] = os.path.basename(root)
                        # Identify Journal from DB map (would be done in next step)
                        all_metadata.append(meta)
                except Exception as e:
                    print(f"Error processing {f}: {e}")

    with open("data/extracted_metadata.json", "w") as out:
        json.dump(all_metadata, out, indent=2)
    print(f"Extraction complete. Processed {len(all_metadata)} files.")

if __name__ == "__main__":
    main()
